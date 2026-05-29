from __future__ import annotations

import asyncio
import logging

from ai_file_brain.core.models import ExtractionResult

logger = logging.getLogger(__name__)


class DocxExtractor:
    async def extract(self, file_path: str) -> ExtractionResult:
        text = await asyncio.to_thread(self._read_sync, file_path)
        return ExtractionResult(text=text, source="native")

    @staticmethod
    def _read_sync(file_path: str) -> str:
        try:
            from docx import Document
        except ImportError as ex:
            logger.warning("python-docx not available; cannot read %s: %s", file_path, ex)
            return ""

        try:
            doc = Document(file_path)
        except FileNotFoundError:
            return ""
        except Exception as ex:
            # Password-protected, malformed, or .doc (legacy) misnamed as .docx.
            logger.warning("Failed to open DOCX %s: %s", file_path, ex)
            return ""

        parts: list[str] = []

        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text:
                            parts.append(p.text)

        for section in doc.sections:
            try:
                header = section.header
            except Exception:
                header = None
            if header is not None:
                for p in header.paragraphs:
                    if p.text:
                        parts.append(p.text)
            try:
                footer = section.footer
            except Exception:
                footer = None
            if footer is not None:
                for p in footer.paragraphs:
                    if p.text:
                        parts.append(p.text)

        return "\n".join(parts)
